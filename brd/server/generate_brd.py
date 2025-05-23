#!/usr/bin/env python3
import sys
import json
import os
import pandas as pd
import base64
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
from PIL import Image
from io import BytesIO
import markdown
import re
from datetime import datetime

def convert_csv_to_table(csv_path):
    """
    Convert a CSV file to a formatted table for the BRD
    """
    try:
        # Read CSV file
        df = pd.read_csv(csv_path)
        
        # Convert to HTML table
        html_table = df.to_html(index=False, classes='table table-bordered')
        
        # Return both HTML table and DataFrame
        return {
            'html': html_table,
            'headers': df.columns.tolist(),
            'rows': df.to_dict('records'),
            'success': True
        }
    except Exception as e:
        print(f"Error converting CSV to table: {str(e)}")
        return {'success': False, 'error': str(e)}

def generate_content_with_ai(prompt, api_key):
    """Generate content using AI with the provided API key."""
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": "gpt-4-turbo-preview",
        "messages": [
            {"role": "system", "content": "You are a professional business document writer specializing in technical Business Requirements Documents."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.7,
        "max_tokens": 2000
    }
    
    try:
        response = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers,
            json=payload
        )
        
        if response.status_code == 200:
            return response.json()["choices"][0]["message"]["content"]
        else:
            print(f"Error: {response.status_code}")
            print(response.text)
            return "An error occurred while generating content."
    except Exception as e:
        print(f"Exception: {e}")
        return "An error occurred while generating content."

def create_table_from_csv(doc, csv_path):
    """Create a formatted table in the document from a CSV file"""
    try:
        # Read CSV file
        df = pd.read_csv(csv_path)
        
        # Create table
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        
        # Add header row with formatting
        hdr_cells = table.rows[0].cells
        for i, column_name in enumerate(df.columns):
            cell = hdr_cells[i]
            cell.text = str(column_name)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
        
        # Add data rows
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)
        
        # Add caption
        caption = doc.add_paragraph("Table: Data from CSV file")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.style = 'Caption'
        
        return True
    except Exception as e:
        doc.add_paragraph(f"Error creating table from CSV: {str(e)}")
        return False

def create_brd_document(data_file, output_file):
    """Create a BRD document from JSON data."""
    # Load data
    with open(data_file, 'r') as f:
        data = json.load(f)
    
    # Extract components
    template = json.loads(data['template'])
    form_data = json.loads(data['formData'])
    business_use_case = data['businessUseCase']
    business_logic = data['businessLogic']
    outputs = json.loads(data['outputs'])
    files = data.get('files', {})
    api_key = data.get('apiKey', '')
    
    # Create document
    doc = Document()
    
    # Document styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Title
    title = doc.add_heading(f"Business Requirements Document: {template['templateName']}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_paragraph.add_run(f"Generated on {datetime.now().strftime('%B %d, %Y')}").italic = True
    
    doc.add_paragraph()  # Spacer
    
    # Overview Section
    doc.add_heading("1. Overview", 1)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Value'
    
    # Add data rows
    for input_field in template['overview']:
        key = input_field['key']
        label = input_field['label']
        value = form_data.get(key, '')
        
        # Handle different types of values
        if isinstance(value, list):
            value = ", ".join(value)
        
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(value)
    
    doc.add_paragraph()  # Spacer
    
    # Technical Section
    doc.add_heading("2. Technical Information", 1)
    
    # Process uploaded files
    if 'csv' in files and os.path.exists(files['csv']):
        doc.add_heading("2.1 Data Table", 2)
        create_table_from_csv(doc, files['csv'])
    
    if 'image' in files and os.path.exists(files['image']):
        doc.add_heading("2.2 System Diagram", 2)
        try:
            doc.add_picture(files['image'], width=Inches(6.0))
            caption = doc.add_paragraph("Figure: System Diagram")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f"Error processing image file: {e}")
    
    if 'doc' in files and os.path.exists(files['doc']):
        doc.add_heading("2.3 Reference Document", 2)
        doc.add_paragraph(f"Reference document: {os.path.basename(files['doc'])}")
    
    doc.add_paragraph()  # Spacer
    
    # Business Use Case Section
    doc.add_heading("3. Business Use Case", 1)
    doc.add_paragraph(business_use_case)
    
    doc.add_paragraph()  # Spacer
    
    # Generate AI content for Business Logic
    doc.add_heading("4. Business Logic", 1)
    doc.add_paragraph(business_logic)
    
    doc.add_paragraph()  # Spacer
    
    # Output Sections
    doc.add_heading("5. Outputs", 1)
    
    for i, output in enumerate(outputs):
        section_title = output['name']
        doc.add_heading(f"5.{i+1} {section_title}", 2)
        
        # Generate content based on output types
        output_types = output['types']
        
        # Generate AI content
        if 'content' in output_types:
            prompt = f"""
            Based on the following information, generate professional content for the '{section_title}' section of a Business Requirements Document:
            
            Business Use Case:
            {business_use_case}
            
            Business Logic:
            {business_logic}
            
            Please provide well-structured, detailed content suitable for a professional BRD. Include appropriate subsections and formatting.
            """
            
            ai_content = generate_content_with_ai(prompt, api_key)
            
            # Convert markdown to rich text
            html_content = markdown.markdown(ai_content)
            
            # Basic HTML to document conversion
            paragraphs = re.split(r'<[/]?p>', html_content)
            for p in paragraphs:
                if p.strip():
                    # Handle headers using a simpler approach
                    for i in range(1, 7):
                        pattern = f'<h{i}>(.*?)</h{i}>'
                        match = re.search(pattern, p)
                        if match:
                            doc.add_heading(match.group(1), i + 2)
                            break
                    else:
                        # Handle lists
                        if '<ul>' in p or '<ol>' in p:
                            list_items = re.findall(r'<li>(.*?)</li>', p)
                            for item in list_items:
                                doc.add_paragraph(item.strip(), style='List Bullet')
                        else:
                            # Strip any remaining HTML tags
                            clean_text = re.sub(r'<.*?>', '', p)
                            if clean_text.strip():
                                doc.add_paragraph(clean_text.strip())
        
        # Table placeholder - if CSV exists, create another table here for the output
        if 'table' in output_types and 'csv' in files and os.path.exists(files['csv']):
            table_heading = doc.add_paragraph()
            table_heading.add_run(f"{section_title} Table Data").bold = True
            create_table_from_csv(doc, files['csv'])
        elif 'table' in output_types:
            doc.add_paragraph("Table data would be included here based on requirements.")
        
        # Image placeholder
        if 'image' in output_types:
            doc.add_paragraph("Visual representation would be included here based on requirements.")
    
    # Save the document
    doc.save(output_file)
    return True

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python generate_brd.py <data_file> <output_file>")
        sys.exit(1)
    
    data_file = sys.argv[1]
    output_file = sys.argv[2]
    
    if create_brd_document(data_file, output_file):
        print(f"BRD document created successfully: {output_file}")
    else:
        print("Failed to create BRD document")
        sys.exit(1) 